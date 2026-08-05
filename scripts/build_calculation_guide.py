from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "calculation-logic-ru.md"
OUTPUT = ROOT / "docs" / "calculation-logic-ru.docx"

INK = RGBColor(31, 45, 61)
ACCENT = RGBColor(94, 126, 129)
ACCENT_DARK = RGBColor(64, 91, 94)
MUTED = RGBColor(101, 112, 123)
SOFT = "EEF5F1"


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after in (
        ("Heading 1", 16, 15, 6),
        ("Heading 2", 13, 11, 4),
        ("Heading 3", 11.5, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT_DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("МОИ ДЕНЬГИ  /  ПАМЯТКА")
    set_font(run, size=9.5, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Как приложение\nсчитает деньги")
    set_font(run, size=26, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Очень простое объяснение: без сложных слов и с примерами")
    set_font(run, size=12.5, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    shade_paragraph(p, SOFT)
    r = p.add_run("Главная мысль: на счете могут лежать деньги, но часть из них уже оставлена на квартиру, еду и другие планы. Только остальное является свободным.")
    set_font(r, size=11.5, bold=True, color=ACCENT_DARK)


def add_inline_markdown(paragraph, text, base_size=11, color=INK):
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part.replace("`", ""))
        set_font(run, size=base_size, bold=index % 2 == 1, color=color)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    add_inline_markdown(p, text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    add_inline_markdown(p, text)


def build():
    doc = Document()
    setup_styles(doc)
    add_title(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()[1:]
    in_short_version = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            heading = line[3:]
            p = doc.add_paragraph(heading, style="Heading 1")
            in_short_version = heading == "Самая короткая версия"
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            continue
        if len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            add_number(doc, line[3:])
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if line.startswith("**") and line.endswith("**"):
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.right_indent = Inches(0.16)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            shade_paragraph(p, SOFT)
            add_inline_markdown(p, line, base_size=11.5, color=ACCENT_DARK)
        else:
            add_inline_markdown(p, line)

    core = doc.core_properties
    core.title = "Как приложение считает деньги"
    core.subject = "Простое описание логики расчетов финансового трекера"
    core.author = "Мои деньги"
    core.keywords = "финансы, расчеты, баланс, свободный остаток"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
